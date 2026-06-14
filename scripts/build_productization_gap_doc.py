from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[3]
OUT = ROOT / "01_專案文件" / "04_開發紀錄" / "NEON_MECHA_ARENA_競品缺口與下一步驗收表_2026-06-11.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Microsoft JhengHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    run.font.size = Pt(9)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        set_cell_text(header_cells[index], header, True)
        set_cell_shading(header_cells[index], "D9EAF7")
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            set_cell_text(cells[index], str(value))
    doc.add_paragraph()
    return table


def add_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = "Microsoft JhengHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
        run.font.color.rgb = RGBColor(20, 37, 63)
    return heading


def add_paragraph(doc, text, bold_label=None):
    paragraph = doc.add_paragraph()
    if bold_label:
      label = paragraph.add_run(bold_label)
      label.bold = True
      label.font.name = "Microsoft JhengHei"
      label._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    run = paragraph.add_run(text)
    run.font.name = "Microsoft JhengHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    run.font.size = Pt(10.5)
    return paragraph


def add_bullets(doc, items):
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        run = paragraph.add_run(item)
        run.font.name = "Microsoft JhengHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
        run.font.size = Pt(10)


def build_doc():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.6)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.7)
    section.right_margin = Cm(1.7)

    styles = doc.styles
    styles["Normal"].font.name = "Microsoft JhengHei"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    styles["Normal"].font.size = Pt(10.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("NEON MECHA ARENA\n競品缺口與下一步驗收表")
    run.bold = True
    run.font.name = "Microsoft JhengHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(31, 78, 121)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("日期：2026-06-11｜用途：開發對照、客戶說明、直播上線前驗收")
    run.font.name = "Microsoft JhengHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(91, 112, 131)

    add_heading(doc, "一、產品定位", 1)
    add_paragraph(
        doc,
        "NEON MECHA ARENA 是 TikTok LIVE 觀眾互動機甲競技場。觀眾掃描直播 QR code 加入紅藍陣營，化身機甲駕駛員，在直播畫面中即時移動、攻擊、搶 MVP，由主播用後台控場、觸發事件、管理人數與輸出戰報。"
    )
    add_bullets(
        doc,
        [
            "核心不變：多人紅藍 PVP、玩家手機控制、直播畫面觀戰、主播後台控場。",
            "差異化方向：不是單純小遊戲，而是可被包裝、導演、截圖、戰報化的直播節目。",
            "商用目標：讓客戶可以快速上手直播，並能用後台數據證明活動熱度與互動價值。"
        ],
    )

    add_heading(doc, "二、競品缺口轉化表", 1)
    add_table(
        doc,
        ["競品方向", "市場常見能力", "可淬取缺口", "本專案落地方式"],
        [
            [
                "TikFinity / Streamer.bot 類工具",
                "禮物、留言、追蹤等事件觸發音效或畫面效果。",
                "互動很強，但多數不是完整遊戲節目；缺少自有戰場、MVP、戰報與玩家身份累積。",
                "保留直播互動入口，但以機甲戰場、駕駛員身份、MVP 與戰報形成節目感。"
            ],
            [
                "Crowd Control 類外掛",
                "觀眾用點數或平台事件影響遊戲。",
                "偏外掛整合，內容主體仍依附別人的遊戲；商用交付難以變成自有 IP。",
                "本專案自建核心玩法與前後台，後續可接禮物事件但不依賴外部遊戲授權。"
            ],
            [
                "TikTok 小遊戲市集 / itch.io 類遊戲",
                "快速可玩、低門檻、畫面模板化。",
                "容易被複製，缺少主播營運工具、數據報表與客戶交付包。",
                "加入主播導演台、事件系統、後台數據中心、素材輸出與 SOP 文件。"
            ],
            [
                "一般機甲競技遊戲",
                "角色、技能、戰鬥動畫、賽季與養成。",
                "畫面精緻但不是直播互動設計，觀眾無法即時加入成為內容。",
                "用機甲視覺語言包裝即時直播互動，保留簡單可懂的紅藍對抗。"
            ],
            [
                "一般直播工具",
                "OBS / TikTok Studio 來源設定、聊天、攝影機與活動目標。",
                "工具本身不會替主播設計一個可賣的互動內容產品。",
                "後台直接提供 TikTok 直式來源、開播 SOP、QR code、口播與社群文案。"
            ],
        ],
    )

    add_heading(doc, "三、競品來源索引", 1)
    add_table(
        doc,
        ["來源", "網址", "觀察重點"],
        [
            [
                "TikFinity",
                "https://tikfinity.zerody.one/",
                "主打 TikTok LIVE 工具、音效、TTS、互動 Overlay、事件與 API，偏直播工具與觸發器。"
            ],
            [
                "Crowd Control",
                "https://crowdcontrol.live/ 與 https://crowdcontrol.live/tiktok/",
                "支援多平台互動效果與遊戲包，TikTok 方向偏禮物、按讚、分享、追蹤觸發遊戲效果。"
            ],
            [
                "StreamToEarn",
                "https://streamtoearn.io/",
                "把 TikTok gifts、likes、follows 轉成 Minecraft、GTA V、Roblox 等遊戲事件，偏外部遊戲整合。"
            ],
            [
                "TikPlay",
                "https://tikplay.games/",
                "主打 TikTok 直播變成互動瀏覽器遊戲，觀眾透過 gifts、likes、comments、follows 控制遊戲。"
            ],
            [
                "Playroom TikTok Integration",
                "https://docs.joinplayroom.com/features/integrations/tiktok",
                "提供 TikTok LIVE 互動遊戲開發整合方向，偏 SDK/開發平台。"
            ],
            [
                "itch.io TikTok LIVE games",
                "https://itch.io/games/tag-tiktok-live",
                "大量低價或模板化 TikTok LIVE 遊戲，顯示市場已有供給，也代表視覺與後台差異化很重要。"
            ],
        ],
    )

    add_heading(doc, "四、本次產品化已執行重點", 1)
    add_table(
        doc,
        ["項目", "本次調整", "驗收方式"],
        [
            [
                "品牌統一",
                "首頁、玩家頁、Spectator、後台、README 統一為 NEON MECHA ARENA。",
                "搜尋舊名稱，確認主要入口不再顯示 TikTok LIVE PVP MVP。"
            ],
            [
                "後台來源流程",
                "後台加入玩家連結、觀戰連結、TikTok 直式來源的複製按鈕。",
                "建立 Session 後按複製，應取得完整 URL。"
            ],
            [
                "TikTok Studio SOP",
                "後台新增最少步驟版開播流程、複製開播步驟、開啟直式預覽。",
                "主播依序建立場次、複製直式來源、加入 Studio 或改用視窗擷取。"
            ],
            [
                "測試階段 Admin Token",
                "測試 token 與後端預設值統一，暫時維持可直接登入。",
                "開 `/admin` 後不需猜 token；正式交付前再改為清空與隱藏。"
            ],
            [
                "商用數據中心",
                "保留玩家上限、觀戰來源、紅藍比例、平均延遲、加入率、互動高峰、操作次數與 MVP。",
                "完成一局後檢查數據卡、回合紀錄、錯誤表與素材輸出。"
            ],
        ],
    )

    add_heading(doc, "五、直播上線前驗收表", 1)
    add_table(
        doc,
        ["驗收項", "操作步驟", "通過標準", "備註"],
        [
            [
                "本機後台",
                "開啟 `/admin`，確認 online，按登入。",
                "可登入、可建立 Session、目前場次顯示 ID。",
                "測試階段 token 預填；正式前需改。"
            ],
            [
                "手機加入",
                "手機與筆電同網段，掃玩家 QR code 或開玩家連結。",
                "手機可進入控制器，後台玩家數增加。",
                "不要用 localhost 給手機。"
            ],
            [
                "第二玩家",
                "另一個分頁或手機再開玩家連結並輸入暱稱。",
                "紅藍雙方各有玩家，後台顯示 2 人。",
                "若滿房會排隊。"
            ],
            [
                "開局與對戰",
                "後台按開始，玩家使用搖桿和攻擊鍵操作。",
                "Spectator 與玩家小視窗同步顯示 HP、命中、勝負。",
                "控制器設定模式需能切左右手、大小與位置。"
            ],
            [
                "戰場事件",
                "後台觸發能量風暴、弱隊護盾、補給投放等事件。",
                "直播畫面有提示，後台事件 Timeline 有紀錄。",
                "Rate limited 通常代表玩家輸入太密，不是系統故障。"
            ],
            [
                "TikTok Studio",
                "複製 TikTok 直式來源到 Studio；若不接受，開直式預覽並用視窗擷取。",
                "直播畫面只看到戰場，不看到 Chrome 分頁、網址列、admin。",
                "正式直播前至少試播一次。"
            ],
            [
                "素材輸出",
                "完成一局後查看摘要、口播、MVP、社群文案。",
                "可複製內容，語意符合本局結果。",
                "可用於直播後貼文。"
            ],
        ],
    )

    add_heading(doc, "六、法律與平台規範提醒", 1)
    add_bullets(
        doc,
        [
            "避免使用現金獎勵、賭博、下注、抽獎保證中獎等語意；本專案定位為娛樂互動與直播效果。",
            "若未來接 TikTok 禮物或付費互動，不應直接讓付費決定勝負，需先確認平台政策與所在地法規。",
            "不要收集不必要個資；目前暱稱、連線狀態、操作紀錄以活動營運為目的。",
            "音效、字體、圖片、機甲素材需使用自有、授權或可商用資源，避免侵權。",
            "向客戶銷售時需明確列出：可交付範圍、直播平台限制、網路需求、維護方式、資料保留時間與不保證流量收益。"
        ],
    )

    add_heading(doc, "七、下一步優先順序", 1)
    add_table(
        doc,
        ["優先序", "工作", "原因", "完成定義"],
        [
            ["P0", "完整跑一次本機雙玩家 + Spectator + Admin 測試", "避免畫面改動後破壞核心遊戲流程", "可加入、開局、攻擊、結算、下一局。"],
            ["P0", "TikTok Studio 實機試播前檢查", "Studio 對本機 HTTP 支援不穩，需確認實際可播方式", "畫面只顯示戰場，不顯示瀏覽器 UI。"],
            ["P1", "整理正式交付 token 與公開部署設定", "目前為測試便利模式，正式不可直接顯示 token", "正式環境 token 隱藏，PUBLIC_ORIGIN 使用 HTTPS 網域。"],
            ["P1", "把競品差異化整理成銷售簡報", "讓客戶看懂為何不是一般小遊戲", "一頁定位、一頁流程、一頁後台、一頁案例。"],
            ["P2", "加入音效與更完整戰鬥動效", "提升直播記憶點與可分享性", "攻擊、命中、勝利、事件皆有一致音效與動畫。"],
        ],
    )

    add_heading(doc, "八、可直接使用的介紹短文", 1)
    add_paragraph(
        doc,
        "NEON MECHA ARENA 是專為 TikTok LIVE 打造的觀眾互動機甲競技場。觀眾不用下載 App，只要掃描直播畫面 QR code，就能加入紅隊或藍隊，使用手機控制器操控機甲移動與攻擊。主播可在後台開局、鎖房、踢人、觸發戰場事件、查看互動數據，並在每局結束後取得 MVP、戰報摘要與社群文案。它不是單純的小遊戲，而是一套可直播、可控場、可交付、可包裝成活動的互動節目系統。"
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_doc()
